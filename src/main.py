# This is ...

# Press Maiusc+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.


# REQUIRED LIBRARIES
'''
We’ll need three libraries for this project. We use pandas to read data from an Excel file,
python-docx for automating .docx (e.g. MS Word, Google docs, etc) file
pywin32 for interacting with Windows APIs. '''

# pip install pandas python-docx pywin32
# pip install python-docx


from docx import Document
from tables import *
from varname import nameof
from utilities import docx_to_pdf
import os
import win32com.client
from io import BytesIO
import requests
from docx import Document
from docx.shared import Inches

''' The following function takes as parameter the list of column names of each single table'''

for i in range(len(sheet)):
    globals()[f'{sheet[i]}'] = extracted_tables[i]

# ignore error at this stage
tables_list = [box_tossicologico,
      dato_EU_commission,
      nomenclatura_botanica,
      dati_mercato,
      dato_bio_farmacologico,
      dato_etnobotanico,
      WHO_OMS,
      REACH,
      procedure_controllo_qualita,
      principi_attivi_markers,
      Ph_Eur,
      campione_riferimento,
      min_sal_ita,
      ISO,
      FUI,
      ESCOP,
      EMA_HMPC,
      EFSA,
      EFSA_2,
      KOME,
      IARC,
      HCM]

# print(nomenclatura_botanica.columns)
print('TABELLE ESTRATTE:\n', sheet)

def make_plant_card():
    '''
    Automate word Document (.docx) with Python-docx and pywin32.
    :return:
    '''

    proc = 'Process 2'
    print(f'{str(proc)}: creating cards...')

    table = nomenclatura_botanica
    id_list = table.ID_pianta.unique()

    for id in id_list:

        document = Document() # to create a .docx file

        # LOGO
        # document.add_picture('brand_logo.png', width=Inches(1)) # todo

        # TITOLO

        table_temp = table[table.ID_pianta == id]
        plant_name = table_temp['nome_scientifico']
        # plant_title = f'{table_temp.ID_pianta.values}: {plant_name.values}'
        plant_title = plant_name.values
        document.add_heading(plant_title, 0)
        p = document.add_paragraph('')
        p.add_run('\n')

        # INSERIMENTO FOTO
        # document.add_picture('id.png', width=Inches(1)) # todo
        # didascalia # todo: campione di riferimento → didascalia

        # INFO DA TABELLA 'nomenclatura_botanica'

        p.add_run('RAW MATERIAL’S COMMON NAME / NOME COMUNE MATERIA GREZZA:').bold = True
        p.add_run('\n')
        p.add_run(table_temp['nome_comune_inglese'].any()).bold = True
        p.add_run(', ')
        p.add_run(table_temp['nome_comune_italiano'].any()).bold = True
        p.add_run('\n')
        p.add_run(table_temp['nome_comune_inglese_url'].any())
        p.add_run('\n')
        p.add_run(table_temp['nome_comune_italiano_url'].any())
        p.add_run('\n\n')

        p.add_run('SCIENTIFIC NAME OF THE PLANT SPECIES/ NOME SCIENTIFICO SPECIE VEGETALE:').bold = True
        p.add_run('\n')
        p.add_run(table_temp['nome_scientifico'].any()).bold = True
        p.add_run('\n')
        p.add_run(table_temp['nome_scientifico_url'].any())
        p.add_run('\n\n')

        p.add_run('SCIENTIFIC SYNONYMS / SINONIMI SCIENTIFICI:').bold = True
        p.add_run('\n')
        p.add_run(table_temp['sin_scientifici'].any()).bold = True
        p.add_run('\n')
        p.add_run(table_temp['sin_scientifici_url'].any())
        p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA 'ISO'
        ###########################

        tab = ISO
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run('GENERAL AND SPECIFIC ISO GUIDELINES RESEARCH/ RICERCA ISO LINEE GUIDA GENERALI E SPECIFICHE:').bold = True
            p.add_run('\n')
            p.add_run(table_temp['ISO_code'].any()).bold = True
            p.add_run(' ')
            p.add_run(table_temp['ISO_title'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['ISO_code_link'].any())
            p.add_run('\n')
            p.add_run(table_temp['ISO_title_link'].any())
            p.add_run('\n')
            p.add_run(table_temp['link_iso'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA 'EMA'
        ###########################

        tab = EMA_HMPC
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:

            p.add_run('RICERCA EMA (MONOGRAFIE, ASSESSMENTS, LINK , REFERENZE):').bold = True

            if ((table_temp['therapeutic_area'].any()) | (table_temp['area_terapeutica'].any())):
                p.add_run('\n')
                p.add_run('Therapeutic area: ').bold = True
                p.add_run(' ')
                p.add_run(table_temp['therapeutic_area'].any()).bold = True
                p.add_run('\n')
                p.add_run('Area terapeutica: ').bold = True
                p.add_run(' ')
                p.add_run(table_temp['area_terapeutica'].any()).bold = True

            p.add_run('\n')
            p.add_run(table_temp['titolo_monografia'].any())
            p.add_run(' (')
            p.add_run(str(table_temp['anno_edizione'].values[0].astype(int)))
            p.add_run(')\n')
            p.add_run(table_temp['link_monografia'].any())

        ###########################
        # INFO DA TABELLE ESCOP,
        # Ph_Eur, FUI, WHO_OMS,
        # KOME, IARC
        ###########################

        tab_list = [ESCOP, Ph_Eur, FUI, WHO_OMS, KOME, IARC]
        tab_names = ['ESCOP', 'Eur Ph', 'FUI', 'WHO', 'KOME', 'IARC']
        p.add_run('\n\nRICERCA MONOGRAFIE (FARMACOPEE, KOMMISSIONE E, WHO, ESCOP, IARC, HEALTH CANADA MONOGRAPH):').bold = True

        for tab in range(len(tab_list)):
            table_temp = tab_list[tab][tab_list[tab].ID_pianta == id]
            if table_temp.empty == False:

                if 'titolo_latino_monogr' in table_temp.columns:
                    p.add_run('\n')
                    p.add_run(tab_names[tab])
                    p.add_run(' (')
                    p.add_run(str(table_temp['anno_edizione'].values[0].astype(int)))
                    p.add_run(') ')
                    p.add_run(table_temp['note'].any())
                    p.add_run(' ')
                    p.add_run(table_temp['titolo_latino_monogr'].any())

                if 'titolo_inglese_monog' in table_temp.columns:
                    p.add_run(', ')
                    p.add_run(table_temp['titolo_inglese_monog'].any())

                if 'pharmaceutical_preparation' in table_temp.columns:
                    p.add_run(', ')
                    p.add_run(table_temp['pharmaceutical_preparation'].any())

                if 'sostanza_attiva' in table_temp.columns:
                    p.add_run(', ')
                    p.add_run(table_temp['sostanza_attiva'].any())
                    p.add_run(', ')
                    p.add_run(table_temp['active_substance'].any())

            p.add_run('\n')
        p.add_run('\n')

        ###########################
        # INFO DA TABELLA 'EFSA_2'
        ###########################

        tab = EFSA_2
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run('RICERCA EFSA – HEALTH CLAIMS AND SAFETY OF USE / CLAIM SALUTISTICI E SICUREZZA D’USO SICUREZZA D’USO:').bold = True
            p.add_run('\n')
            p.add_run(table_temp['titolo_articolo'].any())
            p.add_run('\n EFSA journal (')
            p.add_run(str(table_temp['anno_pubblicazione'].values[0].astype(int)))
            p.add_run('), ')
            if table_temp['note'].any():
                p.add_run(table_temp['note'].any())
                p.add_run(', ')
            p.add_run(table_temp['tipo_pubblicazione'].any())
            p.add_run('\n')
            p.add_run(table_temp['link_articolo'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA 'EFSA'
        ###########################

        tab = EFSA
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run('EFSA COMPENDIUM:').bold = True
            p.add_run('\n')
            p.add_run(table_temp['nome_botanico'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['parti_della_pianta'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['preparazione_sostanza'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['effetti'].any()).bold = True
            p.add_run('\n')

            if table_temp['limiti_massimi'].any():
                p.add_run('Composizione limite / Concern composition: ').bold = True
                p.add_run(table_temp['limiti_massimi'].any()).bold = True
                p.add_run('\n')
            if table_temp['nome_compendium'].any():
                p.add_run(table_temp['nome_compendium'].any())
                p.add_run('\n')
            if table_temp['compendium_link'].any():
                p.add_run(table_temp['compendium_link'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'min_sal_ita'
        ###########################

        tab = min_sal_ita
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run('MINISTRY OF HEALTH RESEARCH/ RICERCA MINISTERO DELLA SALUTE (LISTE POSITIVE/NEGATIVE; POSITIVE AND NEGATIVE LISTS; AVVERTENZE/ WARNINGS; FITOVIGILANZA/ PHYTOVIGILANCE):').bold = True
            p.add_run('\n')
            p.add_run(table_temp['nome_botanico'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['famiglia'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['parti_pianta_tradiz'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['indicazioni_etichetta'].any())
            p.add_run('\n')
            p.add_run(table_temp['altre_indicazioni'].any())
            p.add_run('\n')
            p.add_run(table_temp['referenze_istituzionali'].any())
            p.add_run('\n')
            p.add_run(table_temp['link_decreto'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'principi_attivi_markers'
        ###########################

        tab = principi_attivi_markers
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run(
                'RESEARCH DATA ACTIVE PRINCIPLES / MARKERS:DATO RICERCA PRINCIPI ATTIVI/MARKERS:').bold = True
            p.add_run('\n')
            p.add_run(table_temp['principio_attivo'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['autore'].any())
            p.add_run(', ')
            p.add_run(table_temp['titolo_articolo'].any())
            p.add_run(', ')
            p.add_run(table_temp['rivista'].any())
            p.add_run(', ')
            p.add_run(str(table_temp['anno_pubblicazione'].values[0].astype(int)))
            p.add_run('\n')
            p.add_run(table_temp['link_articolo'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'procedure_controllo_qualita'
        ###########################

        tab = procedure_controllo_qualita
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run(
                'RESEARCH DATA/ OFFICIAL PROCEDURES OF QUALITY CONTROL/ DATO RICERCA/ PROCEDURE UFFICIALI DEL CONTROLLO QUALITA’:').bold = True
            p.add_run('\n')
            p.add_run(table_temp['nome_botanico'].any()).bold = True
            p.add_run(', ')
            p.add_run(table_temp['estratto'].any()).bold = True
            p.add_run(', ')
            p.add_run(table_temp['classi_chimiche'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['autore'].any())
            p.add_run(', ')
            p.add_run(table_temp['titolo_articolo'].any())
            p.add_run(', ')
            p.add_run(table_temp['rivista'].any())
            p.add_run(', ')
            p.add_run(str(table_temp['anno_pubblicazione'].values[0].astype(int)))
            p.add_run('\n')
            p.add_run(table_temp['link_articolo'].any())
            p.add_run('\n')

        # + Eur Ph , FUI
        tab_list = [Ph_Eur, FUI]
        tab_names = ['Eur Ph', 'FUI']
        for tab in range(len(tab_list)):
            table_temp = tab_list[tab][tab_list[tab].ID_pianta == id]
            if table_temp.empty == False:

                if 'titolo_latino_monogr' in table_temp.columns:
                    p.add_run('\n')
                    p.add_run(tab_names[tab])
                    p.add_run(' (')
                    p.add_run(str(table_temp['anno_edizione'].values[0].astype(int)))
                    p.add_run(') ')
                    p.add_run(table_temp['note'].any())
                    p.add_run(' ')
                    p.add_run(table_temp['titolo_latino_monogr'].any())

                if 'titolo_inglese_monog' in table_temp.columns:
                    p.add_run(', ')
                    p.add_run(table_temp['titolo_inglese_monog'].any())

                if 'pharmaceutical_preparation' in table_temp.columns:
                    p.add_run(', ')
                    p.add_run(table_temp['pharmaceutical_preparation'].any())

                if 'sostanza_attiva' in table_temp.columns:
                    p.add_run(', ')
                    p.add_run(table_temp['sostanza_attiva'].any())
                    p.add_run(', ')
                    p.add_run(table_temp['active_substance'].any())

                p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'dato_bio_farmacologico'
        ###########################

        tab = dato_bio_farmacologico
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run(
                'RESEARCH BIO/PHARMACOLOGICAL DATA:DATO RICERCA BIO/FARMACOLOGICO:').bold = True
            p.add_run('\n')
            p.add_run(table_temp['attivita_bio_farmacol'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['autore'].any())
            p.add_run(', ')
            p.add_run(table_temp['titolo_articolo'].any())
            p.add_run(', ')
            p.add_run(table_temp['rivista'].any())
            p.add_run(', ')
            p.add_run(str(table_temp['anno_pubblicazione'].values[0].astype(int)))
            p.add_run('\n')
            p.add_run(table_temp['link_articolo'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'box_tossicologico'
        ###########################

        tab = box_tossicologico
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run(
                'TOXICOLOGICAL DATA/ DATO TOSSICOLOGICO (MONOGRAFIE IARC, LINK E REFERENZE RICERCA):').bold = True
            p.add_run('\n')
            p.add_run(table_temp['effetti'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['autore'].any())
            p.add_run(', ')
            p.add_run(table_temp['titolo_articolo'].any())
            p.add_run(', ')
            p.add_run(table_temp['rivista'].any())
            p.add_run(', ')
            p.add_run(str(table_temp['anno_pubblicazione'].values[0].astype(int)))
            p.add_run('\n')
            p.add_run(table_temp['link_articolo'].any())
            p.add_run('\n\n')

        tab = IARC
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:

            if 'titolo_latino_monogr' in table_temp.columns:
                p.add_run('IARC \n').bold = True
                p.add_run(tab_names[tab])
                p.add_run(' (')
                p.add_run(str(table_temp['anno_edizione'].values[0].astype(int)))
                p.add_run(') ')
                p.add_run(table_temp['note'].any())
                p.add_run(' ')
                p.add_run(table_temp['titolo_latino_monogr'].any())

            if 'titolo_inglese_monog' in table_temp.columns:
                p.add_run(', ')
                p.add_run(table_temp['titolo_inglese_monog'].any())

            if 'pharmaceutical_preparation' in table_temp.columns:
                p.add_run(', ')
                p.add_run(table_temp['pharmaceutical_preparation'].any())

            if 'sostanza_attiva' in table_temp.columns:
                p.add_run(', ')
                p.add_run(table_temp['sostanza_attiva'].any())
                p.add_run(', ')
                p.add_run(table_temp['active_substance'].any())

            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'dato_etnobotanico'
        ###########################

        tab = dato_etnobotanico
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run(
                'ETHNOBOTANICAL DATA/ DATO ETNOBOTANICO:').bold = True
            p.add_run('\n')
            p.add_run(table_temp['parti_della_pianta'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['preparazione'].any())
            p.add_run('\n')
            p.add_run(table_temp['usi_tradizionali'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['autore'].any())
            p.add_run(', ')
            p.add_run(table_temp['titolo_articolo'].any())
            p.add_run(', ')
            p.add_run(table_temp['rivista'].any())
            p.add_run(', ')
            p.add_run(str(table_temp['anno_pubblicazione'].values[0].astype(int)))
            p.add_run('\n')
            p.add_run(table_temp['link_articolo'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'dato_EU_commission'
        ###########################

        tab = dato_EU_commission
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run(
                'DATO EU COMMISSION (FOOD SAFETY/ NOVEL FOOD):').bold = True
            p.add_run('\n')
            p.add_run(table_temp['nome_botanico'].any()).bold = True
            p.add_run('\n')
            p.add_run(table_temp['autore'].any())
            p.add_run(', ')
            p.add_run(table_temp['titolo_articolo'].any())
            p.add_run(', ')
            p.add_run(table_temp['rivista'].any())
            p.add_run(', ')
            p.add_run(str(table_temp['anno_pubblicazione'].values[0].astype(int)))
            p.add_run('\n')
            p.add_run(table_temp['link_articolo'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'dati_mercato'
        ###########################

        tab = dati_mercato
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run(
                'MARKET DATA / DATI MERCATO:').bold = True
            p.add_run('\n')
            p.add_run(table_temp['autore'].any())
            p.add_run(', ')
            p.add_run(table_temp['titolo_articolo'].any())
            p.add_run(', ')
            p.add_run(table_temp['rivista'].any())
            p.add_run(', ')
            p.add_run(str(table_temp['anno_pubblicazione'].values[0].astype(int)))
            p.add_run('\n')
            p.add_run(table_temp['link_articolo'].any())
            p.add_run('\n\n')

        ###########################
        # INFO DA TABELLA
        # 'campione_riferimento'
        ###########################

        tab = campione_riferimento
        table_temp = tab[tab.ID_pianta == id]

        if table_temp.empty == False:
            p.add_run(
                'ORTO E MUSEO BOTANICO DI PISA\nHERBARIUM SPECIMEN/CAMPIONE D’ERBARIO:\n').bold = True
            p.add_run('\n')
            # todo

        '''
            ###########################
            # FOTO ORTO BOTANICO
            ###########################

            if table_temp['link_foto_tavola_bot'].any():
                p.add_run('LIVING PLANT SPECIMEN/CAMPIONE DI PIANTA IN VIVO:\n').bold = True
                p.add_run(table_temp['link_foto_tavola_bot'].any())
                p.add_run('\n')

            if table_temp['didascalia_foto_living_orto'].any():
                p.add_run(table_temp['didascalia_foto_living_orto'].any())
                p.add_run(
                    '\nhttps://uplantdiscover.sma.unipi.it/Home/Details?id=0bf814eb719b417f9d7925a80163717e\n').bold = True

            ###########################
            # FOTO PROGETTO MAPPA
            ###########################
            count = 1

            if table_temp['link_foto_mappa'].any():
                p.add_run('\nPLANT PARTS AND HERBAL DRUGS/ PARTI DI PIANTA E DROGA:\n').bold = True
                p.add_run('\n')


                for link_ in table_temp['link_foto_mappa']:
                    response = requests.get(link_)  # no need to add stream=True
                    # Access the response body as bytes
                    #   then convert it to in-memory binary stream using `BytesIO`
                    binary_img = BytesIO(response.content)
                    # `add_picture` supports image path or stream, we use stream
                    document.add_picture(binary_img, width=Inches(2))
                    p.add_run(f'Fig.{str(count)}')
                    count+=1


            if table_temp['didascalia_mappa'].any():

                for descr in table_temp['didascalia_mappa']:
                    p.add_run(f'Fig.{str(count)}')
                    p.add_run(descr)
                    p.add_run('\nspecimen linked to the MAPPA project PSD_2021/22_UNIPI\n').bold = True
                    count+=1
        '''

        ###########################
        # salva ed esporta il
        # documento
        ###########################

        # salva la scheda in word
        print(f'Exporting and saving card number {id} as Word document')
        document.save(f'../export/doc/{id}.docx')

        # converte e salva tutte le schede in pdf

        script_dir = os.path.dirname(__file__)  # <-- absolute dir the script is in

        src_rel_path = f'../export/doc/{id}.docx'
        src_abs_file_path = os.path.join(script_dir, src_rel_path)

        dst_rel_path = f'../export/pdf/{id}.pdf'
        dst_abs_file_path = os.path.join(script_dir, dst_rel_path)

        print(f'Exporting and saving card number {id} as PDF document')
        docx_to_pdf(src_abs_file_path, dst_abs_file_path)

    print(f'{str(proc)} successfully terminated.')

make_plant_card()